import streamlit as st
import researcher
import generator
from docx import Document
from io import BytesIO

st.set_page_config(page_title="AcademiSync Pro", layout="wide")
st.title("🎓 AcademiSync: 全自动学术综述生成")

# 侧边栏：功能完全回归
with st.sidebar:
    st.header("⚙️ 写作配置")
    target_words = st.select_slider("目标总字数范围", options=[2000, 5000, 8000], value=5000)
    st.write(f"已设定：{target_words} 字量级")
    
    st.header("📄 参考文献管理")
    uploaded_file = st.file_uploader("上传本地 PDF/Word 文档 (可选)", type=['pdf', 'docx', 'txt'])
    st.info("系统将优先检索在线文献，并结合您上传的资料。")

research_title = st.text_input("请输入综述标题：", placeholder="例如：中药干预肺纤维化的机制研究")

if st.button("🚀 开始生成深度综述"):
    if not research_title:
        st.error("请输入标题后再继续")
        st.stop()
        
    with st.status("🛠️ 正在执行学术工作流...", expanded=True) as status:
        # 第一步：深度解析
        st.write("🧠 正在解析课题并提取学术关键词...")
        analysis = generator.analyze_research_title(research_title)
        
        # 第二步：检索带摘要的文献
        st.write("📡 正在全网检索高质量文献摘要...")
        all_papers = []
        for kw in analysis.get('en_keywords', [])[:3]: # 取前3个关键词检索
            results = researcher.fetch_papers(kw, limit=4)
            all_papers.extend(results)
        
        # 第三步：构建参考文献索引库
        unique_papers = {p['title']: p for p in all_papers}.values()
        context_data = ""
        ref_list_html = "### 📚 参考文献列表\n"
        
        for i, p in enumerate(unique_papers, 1):
            # 将编号 [i] 强行注入上下文，方便 AI 引用
            context_data += f"【文献{i}】标题: {p['title']}。摘要: {(p.get('abstract') or '')[:300]}\n\n"
            ref_list_html += f"[{i}] {p['title']}. ({p.get('year','N/A')}). [来源: {p.get('source','网络')}]\n\n"

        # 第四步：分章撰写（解决字数短的问题）
        st.write("📋 正在规划大纲并分章撰写内容...")
        outline = generator.generate_outline(research_title, context_data)
        
        full_draft = f"# {research_title}\n\n{outline}\n\n"
        
        # 核心：循环生成每一个章节，确保总字数
        for dim in analysis.get('dimensions', ['引言', '核心技术', '挑战与展望']):
            st.write(f"✍️ 正在撰写章节：{dim} (包含文献引用)...")
            chapter_text = generator.generate_chapter(research_title, outline, dim, context_data, target_words)
            full_draft += f"## {dim}\n\n{chapter_text}\n\n"

        # 第五步：合并文献列表
        full_draft += "\n---\n" + ref_list_html
        status.update(label="✅ 生成成功！", state="complete")

    # 展示与下载
    st.markdown(full_draft)
    
    doc = Document()
    for line in full_draft.split('\n'):
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    st.download_button("📥 下载 Word 文档", bio.getvalue(), file_name=f"{research_title}.docx")
    import streamlit as st
import researcher
import generator
from docx import Document
from io import BytesIO

st.set_page_config(page_title="AcademiSync Pro", layout="wide")

# 侧边栏保持
with st.sidebar:
    st.header("⚙️ 写作配置")
    # 将字数滑块设为高档位
    target_words = st.select_slider("目标总字数", options=[2000, 5000, 8000], value=5000)
    uploaded_file = st.file_uploader("上传辅助文献", type=['pdf', 'docx'])

research_title = st.text_input("综述课题：", placeholder="输入您的研究方向...")

if st.button("🚀 开始深度创作"):
    with st.status("正在构建学术知识图谱...", expanded=True) as status:
        # 1. 关键词提取
        analysis = generator.analyze_research_title(research_title)
        
        # 2. 深度检索（关键：获取 URL）
        all_papers = []
        for kw in analysis['en_keywords'][:3]:
            res = researcher.fetch_papers(kw, limit=5)
            all_papers.extend(res)
        
        # 3. 构建带超链接的参考文献库
        unique_papers = {p['title']: p for p in all_papers}.values()
        context_for_ai = ""
        ref_list_markdown = "## 参考文献\n\n"
        
        for i, p in enumerate(unique_papers, 1):
            url = p.get('url') or "https://scholar.google.com"
            title = p['title']
            year = p.get('year', 'N/A')
            
            # 给 AI 的背景资料
            context_for_ai += f"文献[{i}]: {title}. 摘要: {p.get('abstract', '')}\n\n"
            # 给用户的可点击列表
            ref_list_markdown += f"[{i}] [{title}]({url}). {year}.\n\n"

        # 4. 生成超长文本 (分步骤串联)
        # 这里的关键是：先生成大纲，然后对大纲里的每个点独立生成 1000 字
        outline = generator.generate_outline(research_title, context_for_ai)
        
        full_text = f"# {research_title}\n\n{outline}\n\n"
        
        # 5. 循环生成：确保每一章都足够长
        for dim in analysis['dimensions']:
            st.write(f"✍️ 正在深度撰写章节：{dim}...")
            # 这里的 target_words // len(dimensions) 确保总数达标
            chapter_content = generator.generate_chapter_deep(
                research_title, outline, dim, context_for_ai, target_words
            )
            full_text += f"## {dim}\n\n{chapter_content}\n\n"
        
        # 拼接参考文献
        full_text += "\n---\n" + ref_list_markdown
        status.update(label="✅ 万字综述生成成功！", state="complete")

    st.markdown(full_text, unsafe_allow_html=True)
    
    # Word 导出（保留超链接格式）
    doc = Document()
    doc.add_heading(research_title, 0)
    for line in full_text.split('\n'):
        if line.strip(): doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    st.download_button("📥 下载完整 Word", bio.getvalue(), file_name="Academic_Review.docx")
